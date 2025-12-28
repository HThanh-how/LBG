#!/usr/bin/env python3
"""
Script để import khung chương trình từ file Word vào database
Version Final: Import từng dòng và track progress
"""
import sys
import os
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from sqlalchemy.orm import Session
from core.database import SessionLocal
from models import User, TeachingProgram, ImportProgress
from core.logging_config import get_logger
from datetime import datetime
import re

logger = get_logger(__name__)


def find_subjects_in_document(doc):
    """
    Tìm tất cả các môn học trong document và trả về list
    """
    subjects_found = []
    for para in doc.paragraphs:
        text = para.text.strip()
        # Tìm pattern: "MÔN TOÁN", "MÔN TIẾNG VIỆT", "2. MÔN TIẾNG VIỆT", etc.
        match = re.match(r'^(\d+\.\s*)?MÔN\s+(.+)$', text, re.IGNORECASE)
        if match:
            subject_name = match.group(2).strip()
            # Loại bỏ số đầu nếu có
            subject_name = re.sub(r'^\d+\.\s*', '', subject_name)
            subjects_found.append(subject_name)
    return subjects_found


def parse_lesson_row(lesson_text, tiets_text, start_lesson_index=1, is_tieng_viet=False):
    """
    Parse một dòng bài học từ bảng
    Trả về list các bài: [(lesson_name, lesson_index), ...]
    
    Logic:
    - Nếu "Ôn tập các số đến 100 (tiết 1, 2)" + "2 tiết" 
      → Tạo 2 bản ghi: ("Ôn tập các số đến 100 (Tiết 1)", 1), ("Ôn tập các số đến 100 (Tiết 2)", 2)
    
    - Với TIẾNG VIỆT: "Bài 1: ...\n- Đọc: ...\n- Viết: ..." + "Tiết 1, 2\nTiết 3"
      → Tạo 3 bản ghi: ("Đọc: ... (Tiết 1)", 1), ("Đọc: ... (Tiết 2)", 2), ("Viết: ...", 3)
    """
    lessons = []
    
    # Làm sạch text
    lesson_text = lesson_text.strip()
    tiets_text = tiets_text.strip()
    
    if is_tieng_viet:
        # Xử lý đặc biệt cho TIẾNG VIỆT
        # Parse "Tiết 1, 2" hoặc "Tiết 1\nTiết 2" thành list các nhóm tiết
        # Mỗi nhóm tương ứng với một sub-point
        tiet_groups = []
        if tiets_text:
            lines = tiets_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Tìm "Tiết X, Y, Z" hoặc "Tiết X"
                tiet_match = re.search(r'Tiết\s+(\d+(?:\s*,\s*\d+)*)', line, re.IGNORECASE)
                if tiet_match:
                    # Parse "1, 2" hoặc "1"
                    numbers = [int(n.strip()) for n in tiet_match.group(1).split(',')]
                    tiet_groups.append(numbers)
        
        # Tách các bài trong lesson_text (bỏ qua "Bài X:")
        lines = lesson_text.split('\n')
        lesson_items = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Bỏ qua dòng "Bài X: ..." (chỉ lấy các sub-points)
            if re.match(r'^Bài\s+\d+:', line, re.IGNORECASE):
                continue
            # Chỉ lấy các dòng bắt đầu bằng "-"
            if line.startswith('-'):
                lesson_items.append(line[1:].strip())
        
        # Map mỗi lesson_item với các tiết tương ứng
        current_index = start_lesson_index
        
        for idx, lesson_item in enumerate(lesson_items):
            if not lesson_item:
                continue
            
            # Làm sạch tên bài
            lesson_name = lesson_item.strip()
            
            # Lấy nhóm tiết tương ứng
            if idx < len(tiet_groups):
                tiet_group = tiet_groups[idx]
                num_tiets = len(tiet_group)
                
                # Tạo các bản ghi
                for i in range(num_tiets):
                    if num_tiets > 1:
                        lesson_name_with_tiet = f"{lesson_name} (Tiết {i + 1})"
                    else:
                        lesson_name_with_tiet = lesson_name
                    lessons.append((lesson_name_with_tiet, current_index + i))
                
                current_index += num_tiets
            else:
                # Nếu không có tiết, mặc định 1 tiết
                lessons.append((lesson_name, current_index))
                current_index += 1
        
        return lessons
    
    # Logic cho các môn khác (TOÁN, etc.)
    # Parse số tiết từ cột "Số tiết học"
    # Format: "2 tiết\n1 tiết\n2 tiết" hoặc "2 tiết"
    num_tiets_list = []
    if tiets_text:
        lines = tiets_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Tìm "X tiết"
            match = re.search(r'(\d+)\s+tiết', line, re.IGNORECASE)
            if match:
                num_tiets_list.append(int(match.group(1)))
    
    # Tách các bài trong lesson_text (mỗi dòng bắt đầu bằng "-")
    lines = lesson_text.split('\n')
    lesson_items = []
    
    current_item = None
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('-'):
            if current_item:
                lesson_items.append(current_item)
            current_item = line[1:].strip()
        else:
            if current_item:
                current_item += " " + line
            else:
                current_item = line
    
    if current_item:
        lesson_items.append(current_item)
    
    if not lesson_items:
        lesson_items = [lesson_text.strip()]
    
    # Xử lý từng bài
    current_index = start_lesson_index
    
    for idx, lesson_item in enumerate(lesson_items):
        if not lesson_item.strip():
            continue
        
        # Làm sạch tên bài (bỏ phần tiết)
        lesson_name = re.sub(r'\(tiết\s+\d+.*?\)', '', lesson_item, flags=re.IGNORECASE)
        lesson_name = lesson_name.strip()
        
        if not lesson_name:
            continue
        
        # Xác định số tiết cho bài này
        num_tiets = 1  # Mặc định 1 tiết
        if num_tiets_list and idx < len(num_tiets_list):
            num_tiets = num_tiets_list[idx]
        else:
            # Thử tìm trong tên bài: "(tiết 1, 2, 3)"
            tiet_range_match = re.search(r'\(tiết\s+(\d+).*?(\d+)', lesson_item, re.IGNORECASE)
            if tiet_range_match:
                start_tiet = int(tiet_range_match.group(1))
                end_tiet = int(tiet_range_match.group(2))
                num_tiets = end_tiet - start_tiet + 1
        
        # Tạo num_tiets bản ghi với cùng tên bài, thêm (Tiết X) nếu có nhiều tiết
        for i in range(num_tiets):
            if num_tiets > 1:
                # Nếu có nhiều tiết, thêm (Tiết X) vào tên bài
                lesson_name_with_tiet = f"{lesson_name} (Tiết {i + 1})"
            else:
                # Nếu chỉ có 1 tiết, không thêm
                lesson_name_with_tiet = lesson_name
            lessons.append((lesson_name_with_tiet, current_index + i))
        
        # Tăng index cho bài tiếp theo
        current_index += num_tiets
    
    # Nếu không parse được gì, trả về bài với tên gốc
    if not lessons and lesson_text:
        lesson_name = re.sub(r'\(tiết\s+\d+.*?\)', '', lesson_text, flags=re.IGNORECASE)
        lesson_name = lesson_name.strip()
        if lesson_name:
            num_tiets = num_tiets_list[0] if num_tiets_list else 1
            for i in range(num_tiets):
                if num_tiets > 1:
                    lesson_name_with_tiet = f"{lesson_name} (Tiết {i + 1})"
                else:
                    lesson_name_with_tiet = lesson_name
                lessons.append((lesson_name_with_tiet, start_lesson_index + i))
    
    return lessons


def get_or_create_progress(db: Session, user_id: int, file_path: str) -> ImportProgress:
    """Lấy hoặc tạo progress record"""
    progress = db.query(ImportProgress).filter(
        ImportProgress.user_id == user_id,
        ImportProgress.file_path == file_path
    ).first()
    
    if not progress:
        progress = ImportProgress(
            user_id=user_id,
            file_path=file_path,
            table_index=0,
            row_index=0,
            lesson_counter=0,
            status="in_progress",
            created_at=datetime.now().date(),
            updated_at=datetime.now().date()
        )
        db.add(progress)
        db.commit()
        db.refresh(progress)
        logger.info(f"Tạo progress mới: table={progress.table_index}, row={progress.row_index}")
    else:
        logger.info(f"Tiếp tục từ: table={progress.table_index}, row={progress.row_index}, counter={progress.lesson_counter}")
    
    return progress


def update_progress(db: Session, progress: ImportProgress, table_idx: int, row_idx: int, subject: str, lesson_counter: int):
    """Cập nhật progress"""
    progress.table_index = table_idx
    progress.row_index = row_idx
    progress.subject_name = subject
    progress.lesson_counter = lesson_counter
    progress.updated_at = datetime.now().date()
    db.commit()


def parse_curriculum_docx(file_path: str, user_id: int = None, db: Session = None, import_directly: bool = False):
    """
    Đọc và parse file Word chứa khung chương trình
    Import từng dòng và track progress
    """
    if db is None:
        db = SessionLocal()
        should_close = True
    else:
        should_close = False
    
    try:
        if user_id is None:
            user = db.query(User).first()
            if not user:
                logger.error("Không tìm thấy user trong database")
                return []
            user_id = user.id
        
        # Lấy hoặc tạo progress
        progress = get_or_create_progress(db, user_id, file_path)
        
        doc = Document(file_path)
        programs = []
        
        logger.info(f"Đang đọc file: {file_path}")
        logger.info(f"Tổng số bảng: {len(doc.tables)}")
        
        # Tìm tất cả môn học trong document
        subjects_list = find_subjects_in_document(doc)
        logger.info(f"Tìm thấy {len(subjects_list)} môn học trong document: {', '.join(subjects_list)}")
        
        # Xử lý từng bảng - map với môn học
        table_start_idx = 1 if len(doc.tables) > len(subjects_list) else 0
        
        # lesson_counter chung cho tất cả các bảng (để tiếp tục đếm khi cùng môn)
        global_lesson_counter = {}
        
        # Khôi phục lesson_counter từ progress
        if progress.subject_name:
            global_lesson_counter[progress.subject_name] = progress.lesson_counter
        
        # Kiểm tra bảng 3 đặc biệt - có thể là TOÁN tiếp theo hoặc TIẾNG VIỆT
        if len(doc.tables) > 2:
            table3 = doc.tables[2]
            # Lấy nội dung bảng 3 để kiểm tra
            table3_content = ""
            for row_idx in range(min(5, len(table3.rows))):
                row = table3.rows[row_idx]
                for cell in row.cells:
                    table3_content += " " + cell.text.strip()
            table3_content = table3_content.upper()
            
            # Keywords để detect môn học
            math_keywords = ['PHÉP NHÂN', 'PHÉP CHIA', 'BẢNG NHÂN', 'SỐ HẠNG', 'THỪA SỐ', 'TÍCH', 'TỔNG CÁC SỐ']
            tv_keywords = ['ĐỌC', 'VIẾT', 'CHỮ', 'TẬP ĐỌC', 'TẬP VIẾT', 'VĂN', 'TỪ', 'CÂU', 'BÀI ĐỌC']
            
            has_math = any(kw in table3_content for kw in math_keywords)
            has_tv = any(kw in table3_content for kw in tv_keywords)
            
            if has_math and not has_tv:
                logger.info(f"Bảng 3 được xác định là TOÁN (dựa trên nội dung: Phép nhân, phép chia)")
        
        for table_idx, table in enumerate(doc.tables):
            logger.info(f"\nĐang xử lý bảng {table_idx + 1}/{len(doc.tables)}...")
            
            # Bỏ qua nếu đã xử lý xong bảng này
            if table_idx < progress.table_index:
                logger.info(f"  Bảng {table_idx + 1} đã được xử lý, bỏ qua")
                continue
            
            # Detect môn học từ nội dung bảng (tìm chính xác từ keywords)
            table_content = ""
            for row_idx in range(min(10, len(table.rows))):
                row = table.rows[row_idx]
                for cell in row.cells:
                    table_content += " " + cell.text.strip()
            table_content = table_content.upper()
            
            # Keywords đặc trưng cho từng môn (ưu tiên keywords đặc trưng hơn)
            subjects_keywords = {
                'TOÁN': ['PHÉP NHÂN', 'PHÉP CHIA', 'BẢNG NHÂN', 'SỐ HẠNG', 'THỪA SỐ', 'TÍCH', 'TỔNG CÁC SỐ', 'PHÉP CỘNG', 'PHÉP TRỪ', 'ÔN TẬP CÁC SỐ', 'SỐ BỊ TRỪ', 'SỐ TRỪ', 'HIỆU'],
                'TIẾNG VIỆT': ['BÉ MAI ĐÃ LỚN', 'THỜI GIAN BIỂU', 'VIẾT CHỮ HOA', 'TỪ VÀ CÂU', 'BẢNG CHỮ CÁI', 'NGHE – VIẾT', 'NHÌN - VIẾT', 'KHU VƯỜN TUỔI THƠ', 'CON SUỐI BẢN TÔI', 'CON ĐƯỜNG LÀNG', 'BÊN CỬA SỔ', 'CHUYỆN BỐN MÙA', 'ĐẦM SEN', 'DÀN NHẠC MÙA HÈ', 'MÙA ĐÔNG Ở VÙNG CAO', 'PHÂN BIỆT'],
                'ĐẠO ĐỨC': ['QUÝ TRỌNG THỜI GIAN', 'NHẬN LỖI VÀ SỬA LỖI', 'BẢO QUẢN ĐỒ DÙNG CÁ NHÂN', 'BẢO QUẢN ĐỒ DÙNG GIA ĐÌNH', 'KÍNH TRỌNG THẦY GIÁO', 'YÊU QUÝ BẠN BÈ', 'QUAN TÂM GIÚP ĐỠ BẠN'],
                'TỰ NHIÊN XÃ HỘI': ['CÁC THẾ HỆ TRONG GIA ĐÌNH', 'NGHỀ NGHIỆP CỦA NGƯỜI THÂN', 'PHÒNG TRÁNH NGỘ ĐỘC', 'GIỮ VỆ SINH NHÀ Ở', 'MỘT SỐ SỰ KIỆN Ở TRƯỜNG', 'NGÀY NHÀ GIÁO VIỆT NAM', 'AN TOÀN VÀ GIỮ VỆ SINH'],
                'ÂM NHẠC': ['KHÁM PHÁ:', 'HÁT:', 'ĐỌC NHẠC:', 'NHẠC CỤ:', 'THANH PHÁCH', 'TAMBOURINE', 'NHÀ GA ÂM NHẠC', 'NGHE NHẠC:', 'MÚA VUI', 'ĐÔ – RÊ – MI', 'SON – LA'],
                'HOẠT ĐỘNG TRẢI NGHIỆM': ['SINH HOẠT DƯỚI CỜ', 'HOẠT ĐỘNG GIÁO DỤC THEO CHỦ ĐỀ', 'SINH HOẠT LỚP', 'BẦU CHỌN LỚP TRƯỞNG', 'THAM GIA LỄ KHAI GIẢNG', 'XÂY DỰNG NỘI QUY LỚP HỌC'],
                'THỂ DỤC': ['ĐỘI HÌNH ĐỘI NGŨ', 'CHUYỂN ĐỘI HÌNH', 'HÀNG DỌC', 'HÀNG NGANG', 'VÒNG TRÒN', 'GIẬM CHÂN TẠI CHỖ', 'VỆ SINH CÁ NHÂN', 'ĐẢM BẢO AN TOÀN', 'ĐHĐN']
            }
            
            # Tìm môn học phù hợp nhất dựa trên số keywords match
            # Ưu tiên keywords dài hơn (đặc trưng hơn)
            best_match = None
            best_score = 0
            for subject_name, keywords in subjects_keywords.items():
                score = 0
                for kw in keywords:
                    if kw in table_content:
                        # Keywords dài hơn có trọng số cao hơn
                        score += len(kw) / 10.0
                if score > best_score:
                    best_score = score
                    best_match = subject_name
            
            # Chỉ chấp nhận nếu có score đủ cao (ít nhất 3.0)
            if best_match and best_score >= 3.0:
                subject = best_match
                logger.info(f"  Môn học: {subject} (detected từ nội dung, score: {best_score:.1f})")
            else:
                # Nếu không detect được, map theo thứ tự (fallback)
                subject_idx = table_idx - table_start_idx
                if subject_idx < 0 or subject_idx >= len(subjects_list):
                    logger.warning(f"  Không có môn học tương ứng cho bảng {table_idx + 1}, bỏ qua")
                    continue
                subject = subjects_list[subject_idx]
                logger.info(f"  Môn học: {subject} (map theo thứ tự)")
            
            if not subject:
                logger.warning(f"  Không tìm thấy tên môn học cho bảng {table_idx + 1}, bỏ qua")
                continue
            
            if table_idx != 2:  # Đã log cho bảng 3 rồi
                logger.info(f"  Môn học: {subject}")
            
            # Tìm header row (thường là hàng 1 hoặc 2)
            header_row_idx = None
            week_col = None
            topic_col = None
            lesson_col = None
            tiet_col = None
            
            for row_idx in range(min(3, len(table.rows))):
                row = table.rows[row_idx]
                cells = [cell.text.strip() for cell in row.cells]
                
                # Tìm các cột
                for col_idx, cell_text in enumerate(cells):
                    cell_upper = cell_text.upper()
                    if 'TUẦN' in cell_upper or 'THÁNG' in cell_upper:
                        week_col = col_idx
                    if 'CHỦ ĐỀ' in cell_upper or 'MẠCH' in cell_upper:
                        topic_col = col_idx
                    if 'TÊN BÀI' in cell_upper or 'BÀI HỌC' in cell_upper:
                        lesson_col = col_idx
                    if 'TIẾT' in cell_upper and 'HỌC' in cell_upper:
                        tiet_col = col_idx
                
                if lesson_col is not None:
                    header_row_idx = row_idx
                    break
            
            if header_row_idx is None:
                logger.warning(f"  Không tìm thấy header row trong bảng {table_idx + 1}")
                continue
            
            # Đoán cột nếu chưa tìm thấy
            if lesson_col is None and len(table.rows) > header_row_idx:
                if len(table.rows[header_row_idx + 1].cells) >= 3:
                    lesson_col = 2
            
            if tiet_col is None and len(table.rows) > header_row_idx:
                if len(table.rows[header_row_idx + 1].cells) >= 4:
                    tiet_col = 3
            
            # Đọc dữ liệu từ các hàng sau header
            if subject not in global_lesson_counter:
                global_lesson_counter[subject] = 0
            
            # Bắt đầu từ row đã import (nếu có)
            start_row = header_row_idx + 1
            if table_idx == progress.table_index and progress.row_index > 0:
                start_row = progress.row_index + 1
                logger.info(f"  Tiếp tục từ hàng {start_row}")
            
            for row_idx in range(start_row, len(table.rows)):
                row = table.rows[row_idx]
                cells = [cell.text.strip() for cell in row.cells]
                
                if not any(cells):
                    continue
                
                # Lấy tên bài và số tiết
                lesson_text = ""
                tiets_text = ""
                
                if lesson_col is not None and lesson_col < len(cells):
                    lesson_text = cells[lesson_col]
                
                if tiet_col is not None and tiet_col < len(cells):
                    tiets_text = cells[tiet_col]
                
                if not lesson_text:
                    continue
                
                # Parse bài học với start_lesson_index từ counter
                start_index = global_lesson_counter[subject] + 1
                is_tv = (subject == 'TIẾNG VIỆT')
                parsed_lessons = parse_lesson_row(lesson_text, tiets_text, start_index, is_tieng_viet=is_tv)
                
                # Import từng bài vào database ngay
                for lesson_name, lesson_index in parsed_lessons:
                    if not lesson_name:
                        continue
                    
                    # Cập nhật counter
                    global_lesson_counter[subject] = max(global_lesson_counter[subject], lesson_index)
                    
                    if import_directly:
                        # Import trực tiếp vào database
                        existing = db.query(TeachingProgram).filter(
                            TeachingProgram.user_id == user_id,
                            TeachingProgram.subject_name == subject,
                            TeachingProgram.lesson_index == lesson_index
                        ).first()
                        
                        if existing:
                            existing.lesson_name = lesson_name
                        else:
                            db_program = TeachingProgram(
                                user_id=user_id,
                                subject_name=subject,
                                lesson_index=lesson_index,
                                lesson_name=lesson_name
                            )
                            db.add(db_program)
                        
                        # Commit sau mỗi bài để track progress
                        db.commit()
                    else:
                        programs.append({
                            "subject_name": subject,
                            "lesson_index": lesson_index,
                            "lesson_name": lesson_name
                        })
                
                # Cập nhật progress sau mỗi hàng
                update_progress(db, progress, table_idx, row_idx, subject, global_lesson_counter[subject])
                logger.info(f"  Đã import hàng {row_idx + 1}: {len(parsed_lessons)} bài, counter={global_lesson_counter[subject]}")
            
            logger.info(f"  Đã parse được {len([p for p in programs if p['subject_name'] == subject])} bài học cho {subject}")
        
        # Đánh dấu hoàn thành
        if import_directly:
            progress.status = "completed"
            progress.updated_at = datetime.now().date()
            db.commit()
            
            total_count = db.query(TeachingProgram).filter(TeachingProgram.user_id == user_id).count()
            logger.info(f"\nImport hoàn thành! Tổng số bài học: {total_count}")
            return []
        else:
            # Sắp xếp và loại bỏ trùng lặp
            programs.sort(key=lambda x: (x["subject_name"], x["lesson_index"]))
            
            seen = set()
            unique_programs = []
            for prog in programs:
                key = (prog["subject_name"], prog["lesson_index"])
                if key not in seen:
                    seen.add(key)
                    unique_programs.append(prog)
                else:
                    existing = next(p for p in unique_programs if (p["subject_name"], p["lesson_index"]) == key)
                    if len(prog["lesson_name"]) > len(existing["lesson_name"]):
                        existing["lesson_name"] = prog["lesson_name"]
            
            logger.info(f"\nTổng kết: Đã parse được {len(unique_programs)} bài học từ {len(set(p['subject_name'] for p in unique_programs))} môn học")
            
            subjects = sorted(set(p['subject_name'] for p in unique_programs))
            logger.info(f"Các môn học: {', '.join(subjects)}")
            
            return unique_programs
    
    except Exception as e:
        if db:
            progress.status = "error"
            progress.updated_at = datetime.now().date()
            db.commit()
        logger.error(f"Lỗi khi import: {str(e)}", exc_info=True)
        raise
    finally:
        if should_close and db:
            db.close()


def import_to_database(programs: list, user_id: int = None, db: Session = None):
    """Import danh sách bài học vào database (legacy - không dùng nữa)"""
    if db is None:
        db = SessionLocal()
        should_close = True
    else:
        should_close = False
    
    try:
        if user_id is None:
            user = db.query(User).first()
            if not user:
                logger.error("Không tìm thấy user trong database")
                return
            user_id = user.id
        else:
            user = db.query(User).filter(User.id == user_id).first()
            if not user:
                logger.error(f"Không tìm thấy user với ID: {user_id}")
                return
        
        logger.info(f"Đang import cho user: {user.username} (ID: {user_id})")
        
        # Xóa các chương trình cũ
        existing_count = db.query(TeachingProgram).filter(TeachingProgram.user_id == user_id).delete()
        logger.info(f"Đã xóa {existing_count} chương trình cũ")
        
        # Tạo các chương trình mới
        db_programs = []
        for prog in programs:
            db_program = TeachingProgram(
                user_id=user_id,
                subject_name=prog["subject_name"],
                lesson_index=prog["lesson_index"],
                lesson_name=prog["lesson_name"]
            )
            db_programs.append(db_program)
        
        if db_programs:
            db.bulk_save_objects(db_programs)
            logger.info(f"Đã tạo {len(db_programs)} chương trình mới")
        
        db.commit()
        
        # Thống kê
        total_count = db.query(TeachingProgram).filter(TeachingProgram.user_id == user_id).count()
        subjects = db.query(TeachingProgram.subject_name).filter(
            TeachingProgram.user_id == user_id
        ).distinct().all()
        
        logger.info(f"Tổng số bài học trong database: {total_count}")
        logger.info(f"Số môn học: {len(subjects)}")
        logger.info("Import thành công!")
        
    except Exception as e:
        db.rollback()
        logger.error(f"Lỗi khi import: {str(e)}", exc_info=True)
        raise
    finally:
        if should_close:
            db.close()


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Import khung chương trình từ file Word")
    parser.add_argument("file_path", help="Đường dẫn đến file Word")
    parser.add_argument("--user-id", type=int, help="ID của user (nếu không có sẽ lấy user đầu tiên)")
    parser.add_argument("--dry-run", action="store_true", help="Chỉ parse, không import vào database")
    parser.add_argument("--reset", action="store_true", help="Reset progress và import lại từ đầu")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.file_path):
        logger.error(f"File không tồn tại: {args.file_path}")
        sys.exit(1)
    
    db = SessionLocal()
    try:
        if args.user_id is None:
            user = db.query(User).first()
            if not user:
                logger.error("Không tìm thấy user trong database")
                sys.exit(1)
            user_id = user.id
        else:
            user_id = args.user_id
        
        # Reset progress nếu cần
        if args.reset:
            progress = db.query(ImportProgress).filter(
                ImportProgress.user_id == user_id,
                ImportProgress.file_path == args.file_path
            ).first()
            if progress:
                db.delete(progress)
                db.commit()
                logger.info("Đã reset progress")
        
        if not args.dry_run:
            # Import trực tiếp từng dòng
            parse_curriculum_docx(args.file_path, user_id=user_id, db=db, import_directly=True)
        else:
            # Chỉ parse, không import
            programs = parse_curriculum_docx(args.file_path, user_id=user_id, db=db, import_directly=False)
            logger.info(f"\nDRY RUN: Đã parse được {len(programs)} bài học")
            for prog in programs[:10]:
                logger.info(f"  {prog['subject_name']} - Tiết {prog['lesson_index']}: {prog['lesson_name'][:50]}")
    finally:
        db.close()

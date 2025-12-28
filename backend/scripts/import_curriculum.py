#!/usr/bin/env python3
"""
Script để import khung chương trình từ file Word vào database
"""
import sys
import os
from pathlib import Path

# Thêm thư mục backend vào path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from sqlalchemy.orm import Session
from core.database import SessionLocal
from models import User, TeachingProgram
from core.logging_config import get_logger
import re

logger = get_logger(__name__)


def parse_curriculum_docx(file_path: str) -> list:
    """
    Đọc và parse file Word chứa khung chương trình
    Trả về danh sách các bài học với format:
    [
        {
            "subject_name": "Tên môn",
            "lesson_index": 1,
            "lesson_name": "Tên bài"
        },
        ...
    ]
    """
    doc = Document(file_path)
    programs = []
    current_subject = None
    
    logger.info(f"Đang đọc file: {file_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Tìm môn học mới (thường là tiêu đề lớn, in đậm, hoặc có pattern đặc biệt)
        # Có thể là: "MÔN: ...", "MÔN HỌC: ...", hoặc chỉ là tên môn đơn giản
        subject_patterns = [
            r'^MÔN[:\s]+(.+)$',
            r'^MÔN HỌC[:\s]+(.+)$',
            r'^([A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ\s]+)$',  # Tất cả chữ in hoa
        ]
        
        # Kiểm tra xem có phải là tiêu đề môn học không
        is_subject = False
        for pattern in subject_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                current_subject = match.group(1).strip() if match.lastindex else text.strip()
                # Làm sạch tên môn
                current_subject = re.sub(r'^MÔN[:\s]+', '', current_subject, flags=re.IGNORECASE)
                current_subject = re.sub(r'^MÔN HỌC[:\s]+', '', current_subject, flags=re.IGNORECASE)
                current_subject = current_subject.strip()
                is_subject = True
                logger.info(f"Tìm thấy môn học: {current_subject}")
                break
        
        if is_subject:
            continue
        
        # Nếu chưa có môn học, bỏ qua
        if not current_subject:
            continue
        
        # Tìm bài học - có thể có format:
        # "Tiết 1: Tên bài"
        # "1. Tên bài"
        # "Bài 1: Tên bài"
        # Hoặc chỉ là số thứ tự + tên bài
        lesson_patterns = [
            r'^Tiết\s+(\d+)[:\s]+(.+)$',
            r'^Bài\s+(\d+)[:\s]+(.+)$',
            r'^(\d+)[\.\)]\s*(.+)$',
            r'^(\d+)[:\s]+(.+)$',
        ]
        
        lesson_index = None
        lesson_name = None
        
        for pattern in lesson_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                try:
                    lesson_index = int(match.group(1))
                    lesson_name = match.group(2).strip()
                    break
                except (ValueError, IndexError):
                    continue
        
        # Nếu không match pattern, có thể là tên bài đơn giản (tiếp theo bài trước)
        if lesson_index is None and lesson_name is None:
            # Kiểm tra xem có phải là số đơn giản không
            if re.match(r'^\d+$', text):
                # Có thể là số tiết, bỏ qua hoặc xử lý đặc biệt
                continue
            # Có thể là tên bài không có số, thêm vào bài trước hoặc tạo bài mới
            if programs and programs[-1]["subject_name"] == current_subject:
                # Nối vào bài trước
                last_program = programs[-1]
                last_program["lesson_name"] += " " + text
                continue
            else:
                # Tạo bài mới với index tự động
                if programs:
                    last_index = max([p["lesson_index"] for p in programs if p["subject_name"] == current_subject], default=0)
                    lesson_index = last_index + 1
                else:
                    lesson_index = 1
                lesson_name = text
        
        if lesson_index and lesson_name:
            programs.append({
                "subject_name": current_subject,
                "lesson_index": lesson_index,
                "lesson_name": lesson_name
            })
            logger.debug(f"  Bài {lesson_index}: {lesson_name}")
    
    # Xử lý bảng (tables) trong Word
    for table in doc.tables:
        current_subject_in_table = None
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            
            # Tìm môn học trong bảng
            for cell_text in cells:
                if not cell_text:
                    continue
                for pattern in subject_patterns:
                    match = re.match(pattern, cell_text, re.IGNORECASE)
                    if match:
                        current_subject_in_table = match.group(1).strip() if match.lastindex else cell_text.strip()
                        current_subject_in_table = re.sub(r'^MÔN[:\s]+', '', current_subject_in_table, flags=re.IGNORECASE)
                        current_subject_in_table = current_subject_in_table.strip()
                        if current_subject_in_table:
                            current_subject = current_subject_in_table
                            logger.info(f"Tìm thấy môn học trong bảng: {current_subject}")
                        break
            
            # Tìm bài học trong bảng
            # Thường có format: [Số tiết, Tên bài] hoặc [Tên bài, Số tiết]
            if current_subject:
                lesson_index = None
                lesson_name = None
                
                # Tìm số tiết
                for cell_text in cells:
                    if re.match(r'^\d+$', cell_text):
                        try:
                            lesson_index = int(cell_text)
                        except ValueError:
                            continue
                        break
                
                # Tìm tên bài (cell không phải số và không rỗng)
                for cell_text in cells:
                    if cell_text and not re.match(r'^\d+$', cell_text):
                        # Kiểm tra xem có phải là tiêu đề cột không
                        if cell_text.upper() in ['TIẾT', 'BÀI', 'TÊN BÀI', 'NỘI DUNG', 'STT']:
                            continue
                        lesson_name = cell_text.strip()
                        break
                
                if lesson_index and lesson_name:
                    programs.append({
                        "subject_name": current_subject,
                        "lesson_index": lesson_index,
                        "lesson_name": lesson_name
                    })
                    logger.debug(f"  Bài {lesson_index}: {lesson_name}")
    
    logger.info(f"Đã parse được {len(programs)} bài học từ {len(set(p['subject_name'] for p in programs))} môn học")
    return programs


def import_to_database(programs: list, user_id: int = None, db: Session = None):
    """
    Import danh sách bài học vào database
    """
    if db is None:
        db = SessionLocal()
        should_close = True
    else:
        should_close = False
    
    try:
        # Lấy user đầu tiên nếu không chỉ định
        if user_id is None:
            user = db.query(User).first()
            if not user:
                logger.error("Không tìm thấy user trong database")
                return
            user_id = user.id
        else:
            user = db.query(User).filter(User.id == user_id).first()
            if not user:
                logger.error(f"Không tìm thấy user với ID: {user_id}")
                return
        
        logger.info(f"Đang import cho user: {user.username} (ID: {user_id})")
        
        # Xóa các chương trình cũ của user (tùy chọn - có thể comment nếu muốn giữ lại)
        # existing_count = db.query(TeachingProgram).filter(TeachingProgram.user_id == user_id).delete()
        # logger.info(f"Đã xóa {existing_count} chương trình cũ")
        
        # Tạo các chương trình mới
        db_programs = []
        for prog in programs:
            # Kiểm tra xem đã tồn tại chưa
            existing = db.query(TeachingProgram).filter(
                TeachingProgram.user_id == user_id,
                TeachingProgram.subject_name == prog["subject_name"],
                TeachingProgram.lesson_index == prog["lesson_index"]
            ).first()
            
            if existing:
                # Cập nhật tên bài nếu khác
                if existing.lesson_name != prog["lesson_name"]:
                    existing.lesson_name = prog["lesson_name"]
                    logger.debug(f"Cập nhật: {prog['subject_name']} - Tiết {prog['lesson_index']}")
            else:
                # Tạo mới
                db_program = TeachingProgram(
                    user_id=user_id,
                    subject_name=prog["subject_name"],
                    lesson_index=prog["lesson_index"],
                    lesson_name=prog["lesson_name"]
                )
                db_programs.append(db_program)
        
        if db_programs:
            db.bulk_save_objects(db_programs)
            logger.info(f"Đã tạo {len(db_programs)} chương trình mới")
        
        db.commit()
        
        # Thống kê
        total_count = db.query(TeachingProgram).filter(TeachingProgram.user_id == user_id).count()
        subjects = db.query(TeachingProgram.subject_name).filter(
            TeachingProgram.user_id == user_id
        ).distinct().all()
        
        logger.info(f"Tổng số bài học trong database: {total_count}")
        logger.info(f"Số môn học: {len(subjects)}")
        logger.info("Import thành công!")
        
    except Exception as e:
        db.rollback()
        logger.error(f"Lỗi khi import: {str(e)}", exc_info=True)
        raise
    finally:
        if should_close:
            db.close()


def main():
    """Main function"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Import khung chương trình từ file Word')
    parser.add_argument('file_path', help='Đường dẫn đến file Word')
    parser.add_argument('--user-id', type=int, help='ID của user (mặc định: user đầu tiên)')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.file_path):
        logger.error(f"File không tồn tại: {args.file_path}")
        return
    
    # Parse file
    programs = parse_curriculum_docx(args.file_path)
    
    if not programs:
        logger.warning("Không tìm thấy dữ liệu trong file")
        return
    
    # Import vào database
    import_to_database(programs, user_id=args.user_id)


if __name__ == "__main__":
    main()


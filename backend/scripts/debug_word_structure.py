#!/usr/bin/env python3
"""
Script để debug cấu trúc file Word
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document

file_path = "../docs/KHUNG CHƯƠNG TRÌNH TRỌN BỘ.25 26.docx"
doc = Document(file_path)

print("=" * 80)
print("CẤU TRÚC FILE WORD")
print("=" * 80)

print(f"\nTổng số paragraphs: {len(doc.paragraphs)}")
print(f"Tổng số tables: {len(doc.tables)}")

# In một số paragraphs đầu
print("\n" + "=" * 80)
print("MỘT SỐ PARAGRAPHS ĐẦU (20 dòng đầu):")
print("=" * 80)
for i, para in enumerate(doc.paragraphs[:20]):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "None"
        print(f"{i+1:3d}. [{style:20s}] {text[:100]}")

# In thông tin về các bảng
print("\n" + "=" * 80)
print("THÔNG TIN VỀ CÁC BẢNG:")
print("=" * 80)
for table_idx, table in enumerate(doc.tables):
    print(f"\n--- BẢNG {table_idx + 1} ---")
    print(f"Số hàng: {len(table.rows)}")
    print(f"Số cột: {len(table.rows[0].cells) if table.rows else 0}")
    
    # In 5 hàng đầu
    print("\n5 hàng đầu tiên:")
    for row_idx in range(min(5, len(table.rows))):
        row = table.rows[row_idx]
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Hàng {row_idx + 1}: {cells}")
    
    # Tìm header row
    print("\nTìm header row:")
    for row_idx in range(min(3, len(table.rows))):
        row = table.rows[row_idx]
        cells = [cell.text.strip() for cell in row.cells]
        has_header_keywords = any(
            keyword in cell.upper() 
            for cell in cells 
            for keyword in ['TIẾT', 'BÀI', 'TÊN', 'MÔN', 'STT', 'NỘI DUNG']
        )
        if has_header_keywords:
            print(f"  → Hàng {row_idx + 1} có vẻ là header: {cells}")

